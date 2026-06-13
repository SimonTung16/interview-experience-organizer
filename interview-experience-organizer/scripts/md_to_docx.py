#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Convert a Markdown interview review document to a Word .docx document.

Usage:
    python scripts/md_to_docx.py \
        --input interview_outputs/03_grounded_interview_review_document.md \
        --output interview_outputs/04_grounded_interview_review_document.docx

Required dependency:
    pip install python-docx
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import List, Optional

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except Exception as e:
    raise RuntimeError(
        "python-docx is required. Install it with: pip install python-docx"
    ) from e


# =========================
# Basic style helpers
# =========================

def set_run_font(
    run,
    font_name: str = "Microsoft YaHei",
    font_size: Optional[float] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
) -> None:
    """Set font for both Western and East Asian text."""
    run.font.name = font_name

    # Ensure Chinese font works in Word/WPS.
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)

    if font_size is not None:
        run.font.size = Pt(font_size)

    if bold is not None:
        run.bold = bold

    if italic is not None:
        run.italic = italic

    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill: str = "F2F2F2") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_run_font(run, font_name="Microsoft YaHei", font_size=9.5, bold=bold)
    set_paragraph_spacing(paragraph, after=0)


def add_horizontal_rule(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("─" * 50)
    set_run_font(run, font_size=9, color="888888")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=8)


def clean_inline_markdown(text: str) -> str:
    """Remove simple inline Markdown markers for fallback plain text."""
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def add_inline_markdown_paragraph(doc: Document, text: str, style: Optional[str] = None):
    """
    Add a paragraph with minimal inline Markdown support:
    - **bold**
    - `inline code`
    """
    paragraph = doc.add_paragraph(style=style)
    set_paragraph_spacing(paragraph)

    # Split into tokens preserving **bold** and `code`.
    pattern = r"(\*\*.*?\*\*|`.*?`)"
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            content = part[2:-2]
            run = paragraph.add_run(content)
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            content = part[1:-1]
            run = paragraph.add_run(content)
            set_run_font(run, font_name="Consolas", font_size=9, color="333333")
        else:
            run = paragraph.add_run(part)
            set_run_font(run)

    return paragraph


def add_code_block(doc: Document, code: str) -> None:
    """Add a code block using monospace font."""
    if not code.strip():
        return

    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=2, after=8, line=1.0)

    run = paragraph.add_run(code.rstrip())
    set_run_font(run, font_name="Consolas", font_size=9)

    # Add light shading to code paragraph.
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F6F8FA")
    p_pr.append(shd)


# =========================
# Markdown table parser
# =========================

def is_markdown_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def is_markdown_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not is_markdown_table_line(stripped):
        return False

    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    if not cells:
        return False

    return all(re.fullmatch(r":?-{3,}:?", cell or "") is not None for cell in cells)


def parse_table_row(line: str) -> List[str]:
    return [clean_inline_markdown(cell.strip()) for cell in line.strip().strip("|").split("|")]


def add_markdown_table(doc: Document, table_lines: List[str]) -> None:
    """Convert simple Markdown table to Word table."""
    if len(table_lines) < 2:
        for line in table_lines:
            add_inline_markdown_paragraph(doc, line)
        return

    # Remove separator line if present.
    rows = []
    for line in table_lines:
        if is_markdown_table_separator(line):
            continue
        rows.append(parse_table_row(line))

    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (max_cols - len(row)) for row in rows]

    table = doc.add_table(rows=len(normalized_rows), cols=max_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(normalized_rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            is_header = r_idx == 0
            set_cell_text(cell, value, bold=is_header)
            if is_header:
                set_cell_shading(cell, fill="EDEDED")

    doc.add_paragraph()


# =========================
# Document style setup
# =========================

def setup_document_styles(doc: Document) -> None:
    """Configure basic Word document styles."""
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.bold = True

    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.size = Pt(15)
    styles["Heading 3"].font.size = Pt(13)
    styles["Heading 4"].font.size = Pt(11.5)


# =========================
# Markdown to DOCX converter
# =========================

def markdown_to_docx(md_text: str, output_path: Path) -> None:
    doc = Document()
    setup_document_styles(doc)

    lines = md_text.splitlines()

    in_code_block = False
    code_lines: List[str] = []

    table_buffer: List[str] = []

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        # Code fence.
        if stripped.startswith("```"):
            flush_table()

            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Markdown table.
        if is_markdown_table_line(stripped):
            table_buffer.append(stripped)
            continue
        else:
            flush_table()

        # Empty line.
        if not stripped:
            continue

        # Horizontal rule.
        if stripped in {"---", "***", "___"}:
            add_horizontal_rule(doc)
            continue

        # Headings.
        if stripped.startswith("# "):
            paragraph = doc.add_heading(clean_inline_markdown(stripped[2:]), level=1)
            set_paragraph_spacing(paragraph, before=6, after=10)
            continue

        if stripped.startswith("## "):
            paragraph = doc.add_heading(clean_inline_markdown(stripped[3:]), level=2)
            set_paragraph_spacing(paragraph, before=8, after=8)
            continue

        if stripped.startswith("### "):
            paragraph = doc.add_heading(clean_inline_markdown(stripped[4:]), level=3)
            set_paragraph_spacing(paragraph, before=6, after=6)
            continue

        if stripped.startswith("#### "):
            paragraph = doc.add_heading(clean_inline_markdown(stripped[5:]), level=4)
            set_paragraph_spacing(paragraph, before=4, after=4)
            continue

        # Bullet list.
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            paragraph = add_inline_markdown_paragraph(doc, text, style="List Bullet")
            set_paragraph_spacing(paragraph, after=3)
            continue

        # Numbered list.
        if re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped).strip()
            paragraph = add_inline_markdown_paragraph(doc, text, style="List Number")
            set_paragraph_spacing(paragraph, after=3)
            continue

        # Normal paragraph.
        add_inline_markdown_paragraph(doc, stripped)

    # Flush unfinished structures.
    flush_table()

    if in_code_block and code_lines:
        add_code_block(doc, "\n".join(code_lines))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


# =========================
# CLI
# =========================

def main() -> int:
    parser = argparse.ArgumentParser(
        description="Convert Markdown interview review document to Word .docx."
    )

    parser.add_argument(
        "--input",
        required=True,
        help="Input Markdown file path.",
    )

    parser.add_argument(
        "--output",
        required=True,
        help="Output DOCX file path.",
    )

    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()

    if not input_path.exists():
        print(f"[ERROR] Input Markdown file does not exist: {input_path}", file=sys.stderr)
        return 1

    if input_path.suffix.lower() not in {".md", ".markdown", ".txt"}:
        print(
            f"[WARNING] Input file extension is not Markdown-like: {input_path.suffix}",
            file=sys.stderr,
        )

    try:
        md_text = input_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        md_text = input_path.read_text(encoding="utf-8-sig")
    except Exception as e:
        print(f"[ERROR] Failed to read input file: {e}", file=sys.stderr)
        return 1

    try:
        markdown_to_docx(md_text, output_path)
    except Exception as e:
        print(f"[ERROR] Failed to convert Markdown to DOCX: {e}", file=sys.stderr)
        return 1

    print("")
    print("[DONE] DOCX generated.")
    print(f"[INFO] Input Markdown: {input_path}")
    print(f"[INFO] Output DOCX: {output_path}")
    print("")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())