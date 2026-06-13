#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Extract text from interview experience screenshots and document files.

Usage:
    python scripts/extract_interview_materials.py --input "./面经截图" --output interview_outputs

Outputs:
    interview_outputs/01_extracted_raw_text.md
    interview_outputs/extracted_items.json

Supported files:
    - Images: .png .jpg .jpeg .webp .bmp .tif .tiff
    - Text:   .txt .md .markdown
    - Word:   .docx
    - PDF:    .pdf

Notes:
    - Image OCR requires Tesseract OCR installed on the system.
    - For Chinese screenshots, Tesseract should have chi_sim language data installed.
    - PDF text extraction uses PyMuPDF if installed.
    - DOCX text extraction uses python-docx if installed.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import sys
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable, List, Optional, Tuple


# =========================
# Optional dependencies
# =========================

try:
    from PIL import Image, ImageOps, ImageFilter
except Exception:
    Image = None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from tqdm import tqdm
except Exception:
    tqdm = None

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document
except Exception:
    Document = None


# =========================
# File types
# =========================

IMAGE_EXTS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".bmp",
    ".tif",
    ".tiff",
}

TEXT_EXTS = {
    ".txt",
    ".md",
    ".markdown",
}

PDF_EXTS = {
    ".pdf",
}

DOCX_EXTS = {
    ".docx",
}


# =========================
# Data structure
# =========================

@dataclass
class ExtractedItem:
    file_path: str
    file_name: str
    file_type: str
    success: bool
    text: str
    error: Optional[str] = None
    text_hash: Optional[str] = None
    char_count: int = 0


# =========================
# Utility functions
# =========================

def iter_files(input_dir: Path) -> Iterable[Path]:
    """Recursively iterate through files in input_dir."""
    for root, _, files in os.walk(input_dir):
        for name in files:
            path = Path(root) / name
            if path.name.startswith("."):
                continue
            yield path


def is_inside(child: Path, parent: Path) -> bool:
    """Return True if child path is inside parent path."""
    try:
        child_resolved = child.resolve()
        parent_resolved = parent.resolve()
        return parent_resolved == child_resolved or parent_resolved in child_resolved.parents
    except Exception:
        return False


def normalize_text(text: str) -> str:
    """Clean OCR or copied text conservatively."""
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Common social-media noise.
    noise_patterns = [
        r"点赞\s*\d*",
        r"评论\s*\d*",
        r"收藏\s*\d*",
        r"转发\s*\d*",
        r"分享\s*\d*",
        r"展开全部",
        r"收起",
        r"关注",
        r"复制链接",
        r"来自.*?客户端",
        r"发布于.*",
    ]

    for pat in noise_patterns:
        text = re.sub(pat, "", text, flags=re.IGNORECASE)

    lines = []
    for line in text.split("\n"):
        line = line.strip()

        if not line:
            lines.append("")
            continue

        line = re.sub(r"\s{2,}", " ", line)
        lines.append(line)

    cleaned = "\n".join(lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)

    return cleaned.strip()


def hash_text(text: str) -> str:
    """Generate stable hash for extracted text."""
    normalized = re.sub(r"\s+", "", text.lower())
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def build_item(
    path: Path,
    file_type: str,
    success: bool,
    text: str,
    error: Optional[str] = None,
) -> ExtractedItem:
    text = normalize_text(text)

    return ExtractedItem(
        file_path=str(path),
        file_name=path.name,
        file_type=file_type,
        success=success and bool(text),
        text=text,
        error=error if not text else None,
        text_hash=hash_text(text) if text else None,
        char_count=len(text),
    )


# =========================
# Text file extraction
# =========================

def read_text_file(path: Path) -> Tuple[bool, str, Optional[str]]:
    """Read txt / md files with common encodings."""
    encodings = [
        "utf-8",
        "utf-8-sig",
        "gbk",
        "gb18030",
        "big5",
    ]

    last_error = None

    for enc in encodings:
        try:
            return True, path.read_text(encoding=enc), None
        except Exception as e:
            last_error = str(e)

    return False, "", last_error


# =========================
# DOCX extraction
# =========================

def read_docx_file(path: Path) -> Tuple[bool, str, Optional[str]]:
    """Read .docx files if python-docx is installed."""
    if Document is None:
        return False, "", "python-docx is not installed."

    try:
        doc = Document(str(path))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract table text if any.
        for table in doc.tables:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = normalize_text(cell.text)
                    if cell_text:
                        cells.append(cell_text)
                if cells:
                    paragraphs.append(" | ".join(cells))

        text = "\n".join(paragraphs)
        return True, text, None

    except Exception as e:
        return False, "", str(e)


# =========================
# PDF extraction
# =========================

def read_pdf_file(path: Path) -> Tuple[bool, str, Optional[str]]:
    """Read PDF text layer if PyMuPDF is installed."""
    if fitz is None:
        return False, "", "PyMuPDF is not installed. Install with: pip install pymupdf"

    try:
        doc = fitz.open(str(path))
        pages = []

        for page_index, page in enumerate(doc, start=1):
            text = page.get_text("text")
            text = normalize_text(text)
            if text:
                pages.append(f"[Page {page_index}]\n{text}")

        doc.close()

        if not pages:
            return False, "", "No extractable text layer found in PDF."

        return True, "\n\n".join(pages), None

    except Exception as e:
        return False, "", str(e)


# =========================
# OCR extraction
# =========================

def ocr_available() -> Tuple[bool, Optional[str]]:
    """Check whether Tesseract OCR is available."""
    if pytesseract is None:
        return False, "pytesseract is not installed."

    try:
        _ = pytesseract.get_tesseract_version()
        return True, None
    except Exception as e:
        return False, str(e)


def preprocess_image_for_ocr(path: Path):
    """
    Basic preprocessing for screenshot OCR.

    Processing:
    - Convert transparent background to white.
    - Convert to grayscale.
    - Auto contrast.
    - Sharpen.
    - Upscale small screenshots.
    """
    if Image is None:
        raise RuntimeError("Pillow is not installed.")

    img = Image.open(path)

    # Convert transparent image to white background.
    if img.mode in ("RGBA", "LA"):
        background = Image.new("RGB", img.size, "WHITE")
        background.paste(img, mask=img.split()[-1])
        img = background

    img = img.convert("L")
    img = ImageOps.autocontrast(img)
    img = img.filter(ImageFilter.SHARPEN)

    width, height = img.size

    # Upscale small screenshots for better OCR.
    if max(width, height) < 1600:
        img = img.resize((width * 2, height * 2))

    return img


def read_image_ocr(path: Path) -> Tuple[bool, str, Optional[str]]:
    """Run OCR on image file."""
    ok, err = ocr_available()
    if not ok:
        return False, "", err

    try:
        img = preprocess_image_for_ocr(path)

        # Try Chinese + English first, then English fallback.
        lang_candidates = [
            "chi_sim+eng",
            "eng",
        ]

        errors = []

        for lang in lang_candidates:
            try:
                config = "--oem 3 --psm 6"
                text = pytesseract.image_to_string(img, lang=lang, config=config)
                text = normalize_text(text)

                if text:
                    return True, text, None

            except Exception as e:
                errors.append(f"{lang}: {e}")

        if errors:
            return False, "", "; ".join(errors)

        return False, "", "OCR returned empty text."

    except Exception as e:
        return False, "", str(e)


# =========================
# File dispatcher
# =========================

def extract_file(path: Path) -> ExtractedItem:
    """Extract text from one file according to extension."""
    suffix = path.suffix.lower()

    if suffix in TEXT_EXTS:
        success, text, error = read_text_file(path)
        return build_item(path, "text", success, text, error)

    if suffix in IMAGE_EXTS:
        success, text, error = read_image_ocr(path)
        return build_item(path, "image_ocr", success, text, error)

    if suffix in DOCX_EXTS:
        success, text, error = read_docx_file(path)
        return build_item(path, "docx", success, text, error)

    if suffix in PDF_EXTS:
        success, text, error = read_pdf_file(path)
        return build_item(path, "pdf", success, text, error)

    return ExtractedItem(
        file_path=str(path),
        file_name=path.name,
        file_type="unsupported",
        success=False,
        text="",
        error=f"Unsupported extension: {suffix}",
        text_hash=None,
        char_count=0,
    )


def dedupe_items(items: List[ExtractedItem]) -> List[ExtractedItem]:
    """
    Deduplicate identical extracted text.

    This deduplicates at material level only.
    It does not deduplicate questions.
    Question-level deduplication should be handled by the Skill / Codex later.
    """
    seen = set()
    result = []

    for item in items:
        if not item.success or not item.text_hash:
            result.append(item)
            continue

        if item.text_hash in seen:
            item.success = False
            item.error = "Duplicate extracted text."
            item.text = ""
            item.char_count = 0
            result.append(item)
            continue

        seen.add(item.text_hash)
        result.append(item)

    return result


# =========================
# Output builders
# =========================

def build_raw_markdown(items: List[ExtractedItem], input_dir: Path) -> str:
    """Build raw extraction Markdown file."""
    total = len(items)
    success_count = sum(1 for x in items if x.success)
    failed_count = total - success_count

    image_count = sum(1 for x in items if x.file_type == "image_ocr")
    text_count = sum(1 for x in items if x.file_type == "text")
    pdf_count = sum(1 for x in items if x.file_type == "pdf")
    docx_count = sum(1 for x in items if x.file_type == "docx")
    unsupported_count = sum(1 for x in items if x.file_type == "unsupported")

    lines = [
        "# 原始面经材料 OCR / 文本提取结果",
        "",
        "## 提取概览",
        "",
        f"- 输入文件夹：`{input_dir}`",
        f"- 文件总数：{total}",
        f"- 图片 OCR 文件数：{image_count}",
        f"- 文本文件数：{text_count}",
        f"- PDF 文件数：{pdf_count}",
        f"- DOCX 文件数：{docx_count}",
        f"- 不支持文件数：{unsupported_count}",
        f"- 成功提取数：{success_count}",
        f"- 失败或跳过数：{failed_count}",
        "",
        "## 成功提取内容",
        "",
    ]

    idx = 1

    for item in items:
        if not item.success:
            continue

        lines.extend(
            [
                f"### Material {idx}: `{item.file_name}`",
                "",
                f"- 来源路径：`{item.file_path}`",
                f"- 类型：{item.file_type}",
                f"- 字符数：{item.char_count}",
                "",
                "```text",
                item.text.strip(),
                "```",
                "",
            ]
        )

        idx += 1

    lines.extend(
        [
            "## 失败、重复或跳过文件",
            "",
        ]
    )

    for item in items:
        if item.success:
            continue

        lines.extend(
            [
                f"- `{item.file_path}`",
                f"  - 类型：{item.file_type}",
                f"  - 原因：{item.error or 'unknown'}",
            ]
        )

    lines.append("")

    return "\n".join(lines)


def write_outputs(items: List[ExtractedItem], input_dir: Path, output_dir: Path) -> None:
    """Write markdown and json outputs."""
    output_dir.mkdir(parents=True, exist_ok=True)

    raw_md = build_raw_markdown(items, input_dir)

    raw_md_path = output_dir / "01_extracted_raw_text.md"
    json_path = output_dir / "extracted_items.json"

    raw_md_path.write_text(raw_md, encoding="utf-8")

    json_path.write_text(
        json.dumps([asdict(x) for x in items], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


# =========================
# Main
# =========================

def main() -> int:
    parser = argparse.ArgumentParser(
        description="Extract text from interview screenshots and documents."
    )

    parser.add_argument(
        "--input",
        required=True,
        help="Input folder containing screenshots, notes, PDFs, or DOCX files.",
    )

    parser.add_argument(
        "--output",
        default="interview_outputs",
        help="Output folder. Default: interview_outputs",
    )

    parser.add_argument(
        "--no-dedupe",
        action="store_true",
        help="Disable material-level duplicate text filtering.",
    )

    args = parser.parse_args()

    input_dir = Path(args.input).expanduser().resolve()
    output_dir = Path(args.output).expanduser().resolve()

    if not input_dir.exists():
        print(f"[ERROR] Input folder does not exist: {input_dir}", file=sys.stderr)
        return 1

    if not input_dir.is_dir():
        print(f"[ERROR] Input path is not a folder: {input_dir}", file=sys.stderr)
        return 1

    output_dir.mkdir(parents=True, exist_ok=True)

    files = []

    for path in iter_files(input_dir):
        if is_inside(path, output_dir):
            continue
        files.append(path)

    iterator = tqdm(files, desc="Extracting") if tqdm else files

    items: List[ExtractedItem] = []

    for path in iterator:
        item = extract_file(path)
        items.append(item)

    if not args.no_dedupe:
        items = dedupe_items(items)

    write_outputs(items, input_dir, output_dir)

    total = len(items)
    success_count = sum(1 for x in items if x.success)
    failed_count = total - success_count

    image_count = sum(1 for x in items if x.file_type == "image_ocr")
    text_count = sum(1 for x in items if x.file_type == "text")
    pdf_count = sum(1 for x in items if x.file_type == "pdf")
    docx_count = sum(1 for x in items if x.file_type == "docx")
    unsupported_count = sum(1 for x in items if x.file_type == "unsupported")

    print("")
    print("[DONE] Extraction completed.")
    print(f"[INFO] Input folder: {input_dir}")
    print(f"[INFO] Output folder: {output_dir}")
    print(f"[INFO] Total files: {total}")
    print(f"[INFO] Images: {image_count}")
    print(f"[INFO] Text files: {text_count}")
    print(f"[INFO] PDFs: {pdf_count}")
    print(f"[INFO] DOCX files: {docx_count}")
    print(f"[INFO] Unsupported files: {unsupported_count}")
    print(f"[INFO] Successful extractions: {success_count}")
    print(f"[INFO] Failed / skipped / duplicated: {failed_count}")
    print(f"[INFO] Raw markdown: {output_dir / '01_extracted_raw_text.md'}")
    print(f"[INFO] JSON: {output_dir / 'extracted_items.json'}")
    print("")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())